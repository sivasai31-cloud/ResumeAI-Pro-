import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_resume_docx(resume_data):
    doc = Document()
    
    # Page Margins (0.5 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    personal = resume_data.get('personal_info', {})
    if isinstance(personal, str):
        try: personal = json.loads(personal)
        except: personal = {}
        
    summary = resume_data.get('summary', '')
    
    experience = resume_data.get('experience', [])
    if isinstance(experience, str):
        try: experience = json.loads(experience)
        except: experience = []
        
    education = resume_data.get('education', [])
    if isinstance(education, str):
        try: education = json.loads(education)
        except: education = []
        
    projects = resume_data.get('projects', [])
    if isinstance(projects, str):
        try: projects = json.loads(projects)
        except: projects = []
        
    skills = resume_data.get('skills', [])
    if isinstance(skills, str):
        try: skills = json.loads(skills)
        except: skills = []

    # Title / Header
    full_name = personal.get('fullName', 'Candidate Name')
    job_title = personal.get('jobTitle', '')
    
    p_name = doc.add_paragraph()
    r_name = p_name.add_run(full_name)
    r_name.font.size = Pt(22)
    r_name.font.bold = True
    r_name.font.color.rgb = RGBColor(37, 99, 235) # #2563EB
    p_name.paragraph_format.space_after = Pt(2)

    if job_title:
        p_title = doc.add_paragraph()
        r_title = p_title.add_run(job_title)
        r_title.font.size = Pt(13)
        r_title.font.color.rgb = RGBColor(124, 58, 237) # #7C3AED
        p_title.paragraph_format.space_after = Pt(4)

    contact_bits = [b for b in [personal.get('email', ''), personal.get('phone', ''), personal.get('location', ''), personal.get('website', '')] if b]
    if contact_bits:
        p_contact = doc.add_paragraph()
        r_contact = p_contact.add_run(" | ".join(contact_bits))
        r_contact.font.size = Pt(9.5)
        r_contact.font.color.rgb = RGBColor(71, 85, 105)
        p_contact.paragraph_format.space_after = Pt(12)

    def add_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)

    # Summary
    if summary:
        add_heading("Professional Summary")
        p_sum = doc.add_paragraph()
        r_sum = p_sum.add_run(summary)
        r_sum.font.size = Pt(10)
        p_sum.paragraph_format.space_after = Pt(8)

    # Experience
    if experience:
        add_heading("Work Experience")
        for exp in experience:
            p_exp = doc.add_paragraph()
            r_pos = p_exp.add_run(exp.get('position', ''))
            r_pos.bold = True
            r_pos.font.size = Pt(10.5)
            
            r_comp = p_exp.add_run(f" — {exp.get('company', '')} ({exp.get('startDate', '')} - {exp.get('endDate', 'Present')})")
            r_comp.font.size = Pt(10)
            r_comp.font.color.rgb = RGBColor(71, 85, 105)
            p_exp.paragraph_format.space_after = Pt(2)
            
            bullets_raw = exp.get('bullets', '')
            if bullets_raw:
                lines = [l.strip('-• ').strip() for l in bullets_raw.split('\n') if l.strip()]
                for l in lines:
                    p_b = doc.add_paragraph(style='List Bullet')
                    r_b = p_b.add_run(l)
                    r_b.font.size = Pt(9.5)
                    p_b.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Education
    if education:
        add_heading("Education")
        for edu in education:
            p_edu = doc.add_paragraph()
            r_deg = p_edu.add_run(edu.get('degree', ''))
            r_deg.bold = True
            r_deg.font.size = Pt(10.5)
            
            r_inst = p_edu.add_run(f", {edu.get('institution', '')} ({edu.get('startDate', '')} - {edu.get('endDate', '')})")
            r_inst.font.size = Pt(10)
            p_edu.paragraph_format.space_after = Pt(4)

    # Projects
    if projects:
        add_heading("Projects")
        for proj in projects:
            p_proj = doc.add_paragraph()
            r_pn = p_proj.add_run(proj.get('name', ''))
            r_pn.bold = True
            if proj.get('techStack'):
                p_proj.add_run(f" [{proj.get('techStack')}]").font.color.rgb = RGBColor(124, 58, 237)
            p_proj.paragraph_format.space_after = Pt(2)
            if proj.get('description'):
                p_desc = doc.add_paragraph(style='List Bullet')
                p_desc.add_run(proj.get('description')).font.size = Pt(9.5)
                p_desc.paragraph_format.space_after = Pt(4)

    # Skills
    if skills:
        add_heading("Technical Skills")
        p_sk = doc.add_paragraph()
        if isinstance(skills, list):
            sk_str = ", ".join(skills)
        else:
            sk_str = str(skills)
        r_sk = p_sk.add_run(sk_str)
        r_sk.font.size = Pt(10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
